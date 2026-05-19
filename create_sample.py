# -*- coding: utf-8 -*-
"""
脚本名: create_sample.py
作用: 快速生成一个带有多种典型违反劳动法条款（身份证号、手机号、试用期超期、不交社保、违规设定违约金、竞业限制超期且无补偿）的 docx 合同测试文件。
"""

import os
from docx import Document

def create_contract_sample():
    doc = Document()
    
    # 添加大标题
    doc.add_heading('劳动合同书', 0)
    
    # 双方元数据段落
    doc.add_paragraph('甲方（用人单位）：北京智明网络科技有限公司')
    doc.add_paragraph('乙方（劳动者）：张三    身份证号：110101199001012345    联系电话：13987654321')
    
    # 期限条款
    doc.add_heading('第一条 合同期限与试用期', level=1)
    doc.add_paragraph('本合同为有固定期限劳动合同，合同期限为一年（12个月），自 2026 年 6 月 1 日起至 2027 年 5 月 31 日止。其中前三个月（即 2026 年 6 月 1 日至 2026 年 8 月 31 日）为乙方入职试用期。')
    
    # 社会保险条款 (高风险点1)
    doc.add_heading('第二条 社会保险约定', level=1)
    doc.add_paragraph('鉴于本公司目前处于初创成长阶段，为降低运营与行政成本，试用期（合同前三个月）内甲方不为乙方缴纳任何社会保险。待乙方转正并签订转正补充协议后，甲方再行统一为乙方办理并补缴社会保险缴费手续。')
    
    # 离职违约金条款 (高风险点2)
    doc.add_heading('第三条 违约金责任', level=1)
    doc.add_paragraph('乙方在劳动合同期限内单方面提出提前离职，或者因个人违纪旷工等原因被解聘的，需一次性向甲方缴纳人民币 5000 元整的违约赔偿金，以补偿公司在日常招聘推广与岗前短期实操培训中产生的高昂综合成本。')
    
    # 竞业限制条款 (中风险点3)
    doc.add_heading('第四条 竞业限制与商业保密', level=1)
    doc.add_paragraph('劳动合同终止或解除后，乙方在三年内，不得进入任何与甲方生产或者经营同类产品、从事同类业务的有竞争关系的其他用人单位工作，也不得自己开业生产或者经营同类产品、业务。在此竞业限制期间，甲方无须向乙方支付任何额外的月度经济补偿金。')
    
    # 保存文档
    output_path = '测试合同样本(带风险条款).docx'
    doc.save(output_path)
    print(f"[OK] 成功生成测试合同样本文件：{os.path.abspath(output_path)}")

if __name__ == '__main__':
    create_contract_sample()
